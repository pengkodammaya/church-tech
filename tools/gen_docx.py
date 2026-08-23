"""Convert workshop tutorial markdown to styled Microsoft Word (.docx) documents.

Usage:
    uv run --with python-docx python tools/gen_docx.py <input.md> <output.docx>

Supported markdown subset:
    # / ## / ### headings
    paragraphs with **bold** and `inline code`
    - bullets (one indent level: "  - ")
    1. numbered steps (rendered as literal "N." prefixes, no Word auto-numbering)
    - [ ] checkbox items
    ``` fenced code blocks ``` (optional language tag ignored)
    | tables | with |---| separator row
    > blockquotes (used for notes/warnings)
    lines starting with [SCREENSHOT: ...] render as capture placeholders
"""

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

NAVY = RGBColor(0x1E, 0x3A, 0x8A)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
GRAY = RGBColor(0x52, 0x60, 0x6D)
CODE_COLOR = RGBColor(0x9D, 0x2B, 0x0F)
ORANGE = RGBColor(0xB4, 0x53, 0x09)

CODE_FILL = "F2F4F7"
QUOTE_FILL = "FFF7E6"
SCREENSHOT_FILL = "FDF3E7"
HEADER_FILL = "DCE6F5"


def set_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_runs_with_inline(paragraph, text):
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = CODE_COLOR
        else:
            paragraph.add_run(token)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        set_shading(p, CODE_FILL)


def add_table(doc, header_cells, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Table Grid"

    def fill_row(cells, values, bold=False, shading=None):
        for cell, value in zip(cells, values):
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            add_runs_with_inline(p, value.strip())
            for run in p.runs:
                run.font.size = Pt(10)
                if bold:
                    run.bold = True
            if shading:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), shading)
                tc_pr.append(shd)

    fill_row(table.rows[0].cells, header_cells, bold=True, shading=HEADER_FILL)
    for i, row in enumerate(rows):
        fill_row(table.rows[i + 1].cells, row)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line):
    return bool(re.match(r"^\s*\|[\s:\-|]+\|\s*$", line))


def setup_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, NAVY),
        ("Heading 3", 12, BLUE),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = ""
    run = footer_p.add_run("Build Church Tech Workshop  ·  Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer_p._p.append(fld)
    run2 = footer_p.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    footer_p._p.append(fld2)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    return doc


def add_cover(doc, title, subtitle):
    for _ in range(5):
        doc.add_paragraph()
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(24)
    p_run = bar.add_run("▄" * 30)
    p_run.font.color.rgb = BLUE
    p_run.font.size = Pt(6)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(subtitle or "")
    run.font.size = Pt(15)
    run.font.color.rgb = GRAY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    run = meta.add_run(f"Build Church Tech Workshop · {date.today().strftime('%B %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_page_break()


def add_toc(doc):
    heading = doc.add_paragraph()
    run = heading.add_run("Contents")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of contents — press Ctrl+A then F9 in Word to populate."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = p.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)

    note = doc.add_paragraph()
    nrun = note.add_run(
        "(If Word asks to update fields when opening this document, choose Yes.)"
    )
    nrun.font.size = Pt(9)
    nrun.font.italic = True
    nrun.font.color.rgb = GRAY

    doc.add_page_break()


def convert(md_path, docx_path):
    md_path = Path(md_path)
    text = open(md_path, encoding="utf-8").read()

    title = md_path.stem
    subtitle = ""

    fm = re.match(r"^---\s*\ntitle:\s*(.+)\nsubtitle:\s*(.+)\n---\s*\n", text)
    if fm:
        title = fm.group(1).strip().strip('"')
        subtitle = fm.group(2).strip().strip('"')
        text = text[fm.end():]

    doc = setup_document()
    add_cover(doc, title, subtitle)
    add_toc(doc)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        if line.startswith("[SCREENSHOT:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("\u25A1  " + line.strip())
            run.font.color.rgb = ORANGE
            run.font.size = Pt(10)
            run.italic = True
            set_shading(p, SCREENSHOT_FILL)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            heading = doc.add_heading("", level=level)
            add_runs_with_inline(heading, m.group(2))
            for run in heading.runs:
                run.font.name = "Calibri"
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            add_runs_with_inline(p, line[2:])
            for run in p.runs:
                run.italic = True
            set_shading(p, QUOTE_FILL)
            i += 1
            continue

        if line.lstrip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_row(line)
            rows = []
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        m = re.match(r"^(\s*)-\s+\[( |x)\]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            box = "\u2611 " if m.group(2) == "x" else "\u2610 "
            run = p.add_run(box)
            run.font.size = Pt(12)
            add_runs_with_inline(p, m.group(3))
            i += 1
            continue

        m = re.match(r"^(\s*)-\s+(.*)$", line)
        if m:
            style = "List Bullet 2" if len(m.group(1)) >= 2 else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline(p, m.group(2))
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.45)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            num = p.add_run(m.group(1) + ".\t")
            num.bold = True
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(Inches(0.45), WD_TAB_ALIGNMENT.LEFT)
            add_runs_with_inline(p, m.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_inline(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"OK {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
